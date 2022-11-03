# -*- coding: utf-8 -*-
from __future__ import annotations

import aiogram.types as types
from typing import NamedTuple

from docxtpl import DocxTemplate
from .searcher import searcher

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK

import random
import os


class _Compilator(object):
    _FONT_FAMILY = 'Times New Roman'
    _FONT_SIZE = 14

    _INDENT = '         '

    def __init__(self: _Compilator) -> None:
        self._engine = Document
    
    def _normalize_part(self, part: str) -> str:
        return f"{_Compilator._INDENT}{part.replace('=', '').strip()}\n"

    def compile(self, file_path: str, content: str) -> None:
        document = self._engine(file_path)

        paragraph = document.add_paragraph()

        is_begin = True
        for part in content.split('\n'):
            normalized_part = self._normalize_part(part)

            if '==' in part:
                if not is_begin:
                    if 'ссылки' in normalized_part.lower():
                        break
                    if 'примечания' in normalized_part.lower():
                        continue
                    if 'литература' in normalized_part.lower():
                        run.add_break(WD_BREAK.PAGE)
                        normalized_part = normalized_part.lower().replace('литература', 'Использованная литература')
                    else:
                        normalized_part = f'\n{normalized_part}'
                else:
                    is_begin = False

                run = paragraph.add_run(normalized_part)

                if '===' in part:
                    run.font.size = Pt(_Compilator._FONT_SIZE)
                else:
                    run.font.size = Pt(_Compilator._FONT_SIZE + 1)
                
                paragraph.add_run('\n')
                run.font.bold = True
            else:
                run = paragraph.add_run(normalized_part)
                run.font.size = Pt(_Compilator._FONT_SIZE)
                run.font.bold = False

            run.font.name = _Compilator._FONT_FAMILY
        
        document.save(file_path)


class _Docxer(object):
    _TEMPLATES_PATH = './storage/templates'
    _TEMP_PATH = './storage/temp/docs'

    _REPORT_TEMPLATE_NAME = 'report.docx'

    def __init__(self: _Docxer) -> None:
        self._template_engine = DocxTemplate
        self._compilator = _Compilator()
        self._searcher = searcher
    
    def _get_template_path(self, template_name: str) -> str:
        return f'{_Docxer._TEMPLATES_PATH}/{template_name}'
    
    def _create_temp_path(self, file_name: str) -> str:
        return f'{_Docxer._TEMP_PATH}/' \
            f"{file_name.replace(' ', '_')}_" \
            f'{random.randint(100000, 999999)}.docx'

    def _delete_file(self, file_path: str) -> None:
        try:
            os.remove(file_path)
        except (PermissionError, FileNotFoundError):
            pass
    
    async def _send_file(self, msg: types.Message, file_path: str) -> None:
        file = open(file_path, 'rb')
        await msg.reply_document(file)
        file.close()

        self._delete_file(file_path)

    async def send_report(self, msg: types.Message, report_data: ReportData) -> None:
        document = self._template_engine(
            self._get_template_path(_Docxer._REPORT_TEMPLATE_NAME)
        )

        document.render({
            'discipline': report_data.discipline,
            'topic': report_data.topic,
            'student': report_data.student,
            'teacher': report_data.teacher
        })

        file_path = self._create_temp_path(
            f'{report_data.discipline}_{report_data.student}_Реферат'
        )

        document.save(file_path)

        self._compilator.compile(file_path, self._searcher.surf(report_data.topic))
        await self._send_file(msg, file_path)


class ReportData(NamedTuple):
    discipline: str
    topic: str
    student: str
    teacher: str


docxer = _Docxer()
